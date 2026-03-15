# ingestion/parser.py

from pathlib import Path
from typing import Callable, Dict, List, Tuple, Optional
from dataclasses import dataclass

from pypdf import PdfReader
from docx import Document as DocxDocument


@dataclass
class ParsedPage:
    """Represents a single page with its text and page number."""
    page_number: int
    text: str


@dataclass
class ParsedDocument:
    """Structured representation of a parsed document."""
    text: str  # Full concatenated text
    pages: List[ParsedPage]  # Page-by-page breakdown
    metadata: Dict[str, any]  # Document metadata (title, author, etc.)


def parse_txt(file_path: Path) -> ParsedDocument:
    """Parse a text file (treated as single page)."""
    text = file_path.read_text(encoding="utf-8")
    return ParsedDocument(
        text=text,
        pages=[ParsedPage(page_number=1, text=text)],
        metadata={}
    )


def parse_pdf(file_path: Path) -> ParsedDocument:
    """Parse a PDF file, preserving page information and extracting metadata."""
    reader = PdfReader(str(file_path))
    pages: List[ParsedPage] = []
    text_parts: List[str] = []

    # Extract text page by page
    for page_num, page in enumerate(reader.pages, start=1):
        extracted = page.extract_text()
        if extracted:
            pages.append(ParsedPage(page_number=page_num, text=extracted))
            text_parts.append(extracted)

    # Extract PDF metadata
    metadata = {}
    if reader.metadata:
        pdf_meta = reader.metadata
        if pdf_meta.title:
            metadata["title"] = str(pdf_meta.title)
        if pdf_meta.author:
            metadata["author"] = str(pdf_meta.author)
        if pdf_meta.subject:
            metadata["subject"] = str(pdf_meta.subject)
        if pdf_meta.creator:
            metadata["creator"] = str(pdf_meta.creator)
        if pdf_meta.producer:
            metadata["producer"] = str(pdf_meta.producer)
        if pdf_meta.creation_date:
            metadata["creation_date"] = str(pdf_meta.creation_date)
        if pdf_meta.modification_date:
            metadata["modification_date"] = str(pdf_meta.modification_date)

    return ParsedDocument(
        text="\n".join(text_parts),
        pages=pages,
        metadata=metadata
    )


def parse_docx(file_path: Path) -> ParsedDocument:
    """Parse a DOCX file (treated as single page for now)."""
    doc = DocxDocument(str(file_path))
    paragraphs = [para.text for para in doc.paragraphs]
    text = "\n".join(paragraphs)
    
    # Extract DOCX metadata
    metadata = {}
    if doc.core_properties:
        props = doc.core_properties
        if props.title:
            metadata["title"] = props.title
        if props.author:
            metadata["author"] = props.author
        if props.subject:
            metadata["subject"] = props.subject
        if props.created:
            metadata["creation_date"] = str(props.created)
        if props.modified:
            metadata["modification_date"] = str(props.modified)

    return ParsedDocument(
        text=text,
        pages=[ParsedPage(page_number=1, text=text)],
        metadata=metadata
    )


def parse_file(file_path: Path) -> ParsedDocument:
    """
    Detect file type and extract text with page information and metadata.
    Returns a ParsedDocument with structured page data.
    """

    suffix = file_path.suffix.lower()

    parsers: dict[str, Callable[[Path], ParsedDocument]] = {
        ".txt": parse_txt,
        ".pdf": parse_pdf,
        ".docx": parse_docx,
    }

    if suffix not in parsers:
        raise ValueError(f"Unsupported file type: {suffix}")

    parsed_doc = parsers[suffix](file_path)

    if not parsed_doc.text.strip():
        raise ValueError("Parsed file is empty.")

    return parsed_doc